"""
core/parser.py — Extract raw text from PDF and DOCX resumes.
"""
import re
import pdfplumber
from docx import Document


def extract_text(filepath: str, ext: str) -> str:
    """Dispatch to the correct parser based on file extension."""
    if ext == "pdf":
        return _extract_pdf(filepath)
    elif ext == "docx":
        return _extract_docx(filepath)
    return ""


    pages = []
    with pdfplumber.open(filepath) as pdf:
        # Limit to first 5 pages for speed and security
        for page in pdf.pages[:5]:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)
    raw = "\n".join(pages)
    return _clean(raw[:10000]) # Cap at 10k characters


def _extract_docx(filepath: str) -> str:
    """Extract paragraph text from a DOCX file."""
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
            if len(lines) > 500: break # Safety break
            
    # Also grab text from tables (limited)
    for table in doc.tables[:3]:
        for row in table.rows[:10]:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    
    raw = "\n".join(lines)
    return _clean(raw[:10000]) # Cap at 10k characters


def _clean(text: str) -> str:
    """Normalise whitespace and remove non-printable characters."""
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)   # keep printable ASCII
    text = re.sub(r"[ \t]{2,}", " ", text)           # collapse spaces
    text = re.sub(r"\n{3,}", "\n\n", text)            # collapse blank lines
    return text.strip()
