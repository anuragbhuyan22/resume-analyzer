import os
import uuid
import io
from flask import Flask, request, jsonify, render_template, send_file
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from docx import Document
from core.parser import extract_text
from core.analyzer import analyze_resume
from core.matcher import match_jobs
from core.improver import generate_suggestions, rewrite_resume_with_ai
from db.postgres import save_analysis, get_analysis, update_suggestions, init_db

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "dev-secret-key")
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB

ALLOWED_EXTENSIONS = {"pdf", "docx"}
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# Initialize PostgreSQL
init_db(app)


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    if "resume" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["resume"]
    if not file or file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Only PDF and DOCX files are allowed"}), 400

    filename = secure_filename(file.filename)
    unique_id = str(uuid.uuid4())
    ext = filename.rsplit(".", 1)[1].lower()
    saved_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{unique_id}.{ext}")
    file.save(saved_path)

    try:
        # 1. Parse text from resume
        text = extract_text(saved_path, ext)
        if not text or len(text.strip()) < 50:
            return jsonify({"error": "Could not extract text. Please upload a text-based PDF or DOCX."}), 400

        # 2. NLP analysis
        analysis = analyze_resume(text)

        # 3. Job role matching
        matches = match_jobs(text)

        # 4. Improvement suggestions (Rule-based only, no AI yet)
        suggestions = generate_suggestions(text, analysis, matches, include_ai=False)

        # 5. Build full result document
        result = {
            "id": unique_id,
            "filename": filename,
            "score": analysis["score"],
            "score_breakdown": analysis["score_breakdown"],
            "skills": analysis["skills"],
            "sections": analysis["sections"],
            "action_verb_count": analysis["action_verb_count"],
            "quantified_achievements": analysis["quantified_achievements"],
            "job_matches": matches,
            "suggestions": suggestions,
            "word_count": len(text.split()),
            "raw_text": text,
        }

        # 6. Persist to MongoDB
        save_analysis(result)

    except Exception as e:
        print(f"ERROR during analysis: {str(e)}")
        return jsonify({"error": "The server is currently busy. Please try again in 1 minute."}), 500
    finally:
        # Clear memory and files
        import gc
        gc.collect() 
        try:
            if os.path.exists(saved_path):
                os.remove(saved_path)
        except OSError:
            pass

    return jsonify({"id": unique_id})


@app.route("/results/<analysis_id>")
def results(analysis_id):
    return render_template("results.html", analysis_id=analysis_id)


@app.route("/api/results/<analysis_id>")
def api_results(analysis_id):
    data = get_analysis(analysis_id)
    if not data:
        return jsonify({"error": "Analysis not found"}), 404
    return jsonify(data)


@app.route("/api/results/<analysis_id>/ai_suggestions", methods=["POST"])
def get_ai_tips(analysis_id):
    data = get_analysis(analysis_id)
    if not data:
        return jsonify({"error": "Analysis not found"}), 404
    
    from core.improver import _get_gemini_suggestions
    try:
        ai_tips = _get_gemini_suggestions(data["raw_text"], data, data["job_matches"])
        # Update the stored analysis with the new AI tips
        current_suggestions = list(data.get("suggestions", []))
        current_suggestions.extend(ai_tips)
        update_suggestions(analysis_id, current_suggestions)
        return jsonify(ai_tips)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/results/<analysis_id>/suggestions", methods=["POST"])
def update_api_suggestions(analysis_id):
    suggestions = request.json.get("suggestions", [])
    success = update_suggestions(analysis_id, suggestions)
    if success:
        return jsonify({"status": "ok"})
    return jsonify({"error": "Failed to update"}), 500


@app.route("/api/results/<analysis_id>/download_improved", methods=["GET"])
def download_improved(analysis_id):
    data = get_analysis(analysis_id)
    if not data:
        return jsonify({"error": "Analysis not found"}), 404

    raw_text = data.get("raw_text", "")
    suggestions = data.get("suggestions", [])

    if not raw_text:
        return jsonify({"error": "Raw text not available for this resume."}), 400

    try:
        # Generate improved text using AI
        improved_text = rewrite_resume_with_ai(raw_text, suggestions)
        
        # Create DOCX
        doc = Document()
        for paragraph in improved_text.split("\n"):
            # If the paragraph looks like a header (all caps, short), make it a heading
            p_strip = paragraph.strip()
            if not p_strip:
                continue
            if p_strip.isupper() and len(p_strip) < 30:
                doc.add_heading(p_strip, level=1)
            else:
                doc.add_paragraph(p_strip)
        
        # Save to memory
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = f"Improved_{data.get('filename', 'Resume.docx')}"
        if not filename.endswith(".docx"):
            filename = filename.rsplit(".", 1)[0] + ".docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    debug = os.getenv("FLASK_ENV", "development") == "development"
    app.run(host="0.0.0.0", port=port, debug=debug)
