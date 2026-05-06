from docx import Document

doc = Document()
doc.add_heading('John Doe - Software Engineer', 0)
doc.add_paragraph('john.doe@email.com | 555-1234')
doc.add_heading('Summary', level=1)
doc.add_paragraph('I am a software engineer with experience in Python and JavaScript. I made a website.')
doc.add_heading('Experience', level=1)
doc.add_paragraph('Software Developer at Tech Inc.')
doc.add_paragraph('- I wrote code for the backend.')
doc.add_paragraph('- Fixed bugs.')
doc.save('dummy_resume.docx')
print('Created dummy_resume.docx')
